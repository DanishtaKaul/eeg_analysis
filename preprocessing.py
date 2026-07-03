# -*- coding: utf-8 -*-

"""Filter and detrend raw EEG, flag bad channels by kurtosis, then run ICA with ICLabel
and several artifact-detection methods to remove non-brain components"""


import os
import sys
import datetime
import numpy as np
import matplotlib.pyplot as plt
import mne
from mne_icalabel import label_components
from mne.preprocessing import ICA
from mne.io import RawArray
from scipy.stats import kurtosis
from mne.filter import detrend
from docx import Document
from docx.shared import Inches
import pandas as pd


# Set up logging
os.makedirs("logs", exist_ok=True)
log_filename = datetime.datetime.now().strftime(
    "logs/ica_log_%Y-%m-%d_%H-%M-%S.txt")
log_file = open(log_filename, "w", encoding="utf-8")
sys.stdout = type("Tee", (), {
    "write": lambda self, msg: [s.write(msg) and s.flush() for s in (sys.__stdout__, log_file)],
    "flush": lambda self: [s.flush() for s in (sys.__stdout__, log_file)]
})()


def filter_and_detrend_data(raw, kurtosis_z_thresh=5.0):
    # Step 1: Bandpass filter
    raw.filter(l_freq=1.0, h_freq=100.0, fir_design='firwin')

    # Step 2: Detrend 
    raw._data = detrend(raw.get_data(), order=1, axis=1)

    # Step 3: Compute kurtosis
    eeg_data = raw.get_data(picks='eeg')  # (n_channels, n_times)
    kurt_vals = kurtosis(eeg_data, axis=1)
    z_kurt = (kurt_vals - np.mean(kurt_vals)) / np.std(kurt_vals)

    # Step 4: Identify bad channels
    bad_channels = []
    ppid = raw.info['subject_info'].get('pid_str', 'unknown_pid')

    for idx, z in enumerate(z_kurt):
        ch_name = raw.ch_names[idx]
        if z > kurtosis_z_thresh:
            print(f"[Kurtosis] Marking channel {ch_name} (z = {z:.2f}) as bad")
            bad_channels.append((ch_name, z))

    # Step 5: Mark bads
    raw.info['bads'].extend([ch for ch, _ in bad_channels])

    # Step 6: Save per-participant CSV
    os.makedirs("bad_channel_logs", exist_ok=True)
    per_pid_path = os.path.join("bad_channel_logs", f"bad_channels_{ppid}.csv")
    if bad_channels:
        pd.DataFrame(bad_channels, columns=["Channel", "Z_Kurtosis"]).to_csv(
            per_pid_path, index=False)
        print(f"[Kurtosis] Bad channels saved to {per_pid_path}")
    else:
        pd.DataFrame(columns=["Channel", "Z_Kurtosis"]
                     ).to_csv(per_pid_path, index=False)
        print(f"[Kurtosis] No bad channels for {ppid}. Empty log saved.")

    # Step 7: Save to a master CSV across all participants
    log_file = os.path.join("bad_channel_logs", "all_bad_channels.csv")
    log_rows = [(ppid, ch, z) for ch, z in bad_channels]
    if log_rows:
        if os.path.exists(log_file):
            df_existing = pd.read_csv(log_file)
            # Avoid duplicates
            df_existing = df_existing[df_existing['Participant'] != ppid]
            df_all = pd.concat([df_existing, pd.DataFrame(log_rows, columns=[
                               "Participant", "Channel", "Z_Kurtosis"])], ignore_index=True)
        else:
            df_all = pd.DataFrame(
                log_rows, columns=["Participant", "Channel", "Z_Kurtosis"])
        df_all.to_csv(log_file, index=False)
        print(
            f"[Kurtosis] Logged {len(bad_channels)} bad channels to master file.")

    # Step 8: Interpolate bad channels
    if raw.info['bads']:
        raw.interpolate_bads(reset_bads=True)
        print(f"[Kurtosis] Interpolated {len(bad_channels)} bad channels.")

    # Step 9: Rereference
    raw.set_eeg_reference('average', projection=False)

    return raw


def create_report_document(labels, ica, raw, ppid, all_flagged_inds, document_comments):
    os.makedirs("iclabel_reports", exist_ok=True)
    doc = Document()
    doc.add_heading(f"Report for {ppid}", level=1)

    doc.add_heading("ICA Topographies", level=2)
    for i, fig in enumerate(ica.plot_components(show=False)):  # plot ica topography
        fig_path = f"iclabel_reports/{ppid}_component_{i}.png"
        fig.savefig(fig_path, dpi=150)
        plt.close(fig)
        doc.add_picture(fig_path, width=Inches(6))

    doc.add_heading('ICLabel', level=1)  # Add a section header for ICLabel
    # Subheading for the list of components
    doc.add_heading("Component Classification", level=2)

    # Loop through all ICLabel classification results that were saved earlier
    for item in document_comments['ICLabel']:
        # The full comment string for this component (e.g., "Component 01: Label = eye blink, Confidence = 0.98")
        comment = item['comment']
        # Boolean value: True if the component was flagged, False otherwise
        flagged = item['flagged']

        if flagged:
            # If the component was flagged as an artifact, append "(flagged)" to the comment
            doc.add_paragraph(comment + " (flagged)")
        else:
            # Otherwise, just add the comment as-is
            doc.add_paragraph(comment)

    if all_flagged_inds:
        doc.add_heading("Excluded Component Properties (All Methods)", level=2)
        for i, fig in enumerate(ica.plot_properties(raw, picks=list(all_flagged_inds), show=False, psd_args=dict(fmax=50))):

            
            fig_path = f"iclabel_reports/{ppid}_excluded_{i}.png"
            fig.savefig(fig_path, dpi=150)
            plt.close(fig)
            doc.add_picture(fig_path, width=Inches(6))

    document_keys = ['EOG', 'Auto', 'Focality']

    # === Plot properties of retained components for manual inspection ===
    doc.add_heading("Remaining (Retained) Component Properties", level=2)
    retained_inds = [i for i in range(
        ica.n_components_) if i not in all_flagged_inds]

    if retained_inds:
        for i, fig in enumerate(ica.plot_properties(raw, picks=retained_inds, show=False, psd_args=dict(fmax=50))):
            fig_path = f"iclabel_reports/{ppid}_retained_{i}.png"
            fig.savefig(fig_path, dpi=150)
            plt.close(fig)
            doc.add_picture(fig_path, width=Inches(6))
    else:
        doc.add_paragraph(
            "All components were flagged; no retained components to show.")

    for key in document_keys:
        doc.add_heading(key, level=1)

        for item in document_comments[key]:
            comment = item['comment']
            flagged = item['flagged']
            if flagged:
                doc.add_paragraph(comment + " (flagged)")
            else:
                doc.add_paragraph(comment)

    doc.save(f"iclabel_reports/{ppid}_ICLabel_Report.docx")
    print(f"ICLabel report saved for {ppid}")


def apply_ICA(raw):
    """
    Apply ICA to concatenated EEG data and perform multiple artifact detection methods
   
    """

    # Step 0: Prepare raw data
    # Create a full memory-loaded copy of the Raw object
    raw = raw.copy().load_data()
    # Save annotations to apply back to the raw file after ica
    annotations = raw.annotations.copy()

    document_comments = {
        "ICLabel": [],
        "EOG": [],
        "Auto": [],
        "Focality": [],
        
    }

    # Remove annotations that start with 'BAD' (e.g., bad segments or boundary markers)
    bad_annot_inds = [i for i, desc in enumerate(
        raw.annotations.description) if desc.lower().startswith('bad')]
    raw.annotations.delete(bad_annot_inds)  # Delete 'BAD' annotations
    # Delete any remaining annotations for clean ICA
    raw.annotations.delete(np.arange(len(raw.annotations)))

    # Print confirmation of cleanup
    print(f"Removed {len(bad_annot_inds)} BAD annotations for ICA")
    print(f"Raw duration: {raw.n_times} samples")
    print(f"Files used: {raw.filenames}")

    # Retrieve participant ID (used later for manual rules)
    ppid = raw.info['subject_info'].get('pid_str')

    # Step 1: Fit ICA using extended Infomax on raw data wrapped as a RawArray
    raw_data = raw.get_data()  # Extract EEG signal matrix (n_channels x n_times)
    # Create a clean Raw object for ICA, KEEP RANDOM STATE 42 for reproducibility
    raw_clean = RawArray(raw_data, raw.info.copy())
    ica = ICA(n_components=0.999999, random_state=42, method='infomax',
              fit_params=dict(extended=True))
    print("Max value in raw:", np.max(np.abs(raw.get_data())))
    print("Max value in raw_clean:", np.max(np.abs(raw_clean.get_data())))

    # Fit ICA without using annotations
    ica.fit(raw_clean, reject_by_annotation=False)

    # Step 2: Run ICLabel to classify each ICA component
    # Predict labels for ICA components

    labels = label_components(raw, ica, method='iclabel')
    # input(labels)

    artifact_inds = set()
    for i, label in enumerate(labels['labels']):
        y_pred_proba = labels["y_pred_proba"][i]
        if label != 'brain' and y_pred_proba >= 0.7:
            artifact_inds.add(i)

        if label == 'muscle artifact' and y_pred_proba >= 0.5:
            artifact_inds.add(i)
        if label == 'other':
            artifact_inds.add(i)

    # Print each flagged component and its confidence score
    print("\n=== ICLabel Artifact Detection ===")
    for i, label in enumerate(labels['labels']):
        conf = labels['y_pred_proba'][i].max()
        comment = f"Component {i:02d}: Label = {label}, Confidence = {conf:.2f}"
        if i in artifact_inds:
            print(comment + " (Flagged)")
            document_comments['ICLabel'].append(
                {"comment": comment, "flagged": True})
        else:
            print(comment)
            document_comments['ICLabel'].append(
                {"comment": comment, "flagged": False})

    # Step 3: Detect eye-related components using EOG correlation with dynamic z threshold
    z_thresh = 3.5  # Initial z-threshold
    eog_inds = set()
    # Lower threshold if too few components are flagged
    while len(eog_inds) < 3 and z_thresh > 0:
        # Only use EOG channels that are not marked bad

        eog_channels = [ch for ch in ['FP1', 'FP2', 'F7', 'F8']
                        if ch not in raw.info['bads']]

        print(f"EOG channels passed to ICA EOG check: {eog_channels}")

        eog_inds, _ = ica.find_bads_eog(
            raw, ch_name=eog_channels, threshold=z_thresh)

        z_thresh -= 0.05
    eog_inds = set(eog_inds)

    # Print EOG-related components with final threshold used
    print(f"\n=== EOG Correlation (z > {z_thresh + 0.05:.2f}) ===")
    for comp in sorted(eog_inds):
        comment = f"Component {comp:02d}: Flagged by EOG correlation"
        print(comment)
        document_comments['EOG'].append({"comment": comment, "flagged": True})

    # Step 4: Autocorrelation check - flag components with low temporal autocorrelation. 
    def compute_autocorr(sources, lag):
        # For each ICA component src, it shifts the signal forward by lag samples, then it computes the Pearson correlation between the original and the shifted version (lagged autocorrelation)
        return np.array([np.corrcoef(src[:-lag], src[lag:])[0, 1] for src in sources])

    # Timecourses of ICA components (n_components x n_times)
    # (sources): each row is one component, sources contains timecourse of each component over time
    sources = ica.get_sources(raw_clean).get_data()
    lag = int(0.020 * raw.info['sfreq'])  # Lag of 20ms in samples
    autocorrs = compute_autocorr(sources, lag)  # compute autocorrelation
    z_auto = (autocorrs - autocorrs.mean()) / autocorrs.std()  # Z-score
    # Flag components with z < -1.5
    auto_inds = set(np.where(z_auto < -1.5)[0])

    print("\n=== Autocorrelation Check (z < -1.5) ===")
    for i, z in enumerate(z_auto):
        comment = f"Component {i:02d}: Autocorr z = {z:.2f}"
        if i in auto_inds:
            print(comment + " (Flagged)")
            document_comments['Auto'].append(
                {"comment": comment, "flagged": True})
        else:
            print(comment)
            document_comments['Auto'].append(
                {"comment": comment, "flagged": False})

    # Step 5: Focality (Kurtosis of ICA spatial maps). 
    maps = ica.get_components().T  # ICA topographies (n_components x n_channels)
    kurt = kurtosis(maps, axis=1)  # Compute kurtosis for each component
    # Converts each kurtosis value into a z-score which tells how extreme each component's kurtosis is compared to the others
    z_kurt = (kurt - kurt.mean()) / kurt.std()
    focal_inds = set(np.where(z_kurt > 2)[0])  # Flag highly focal components

    print("\n=== Focality (Kurtosis z > 2) ===")
    for i, z in enumerate(z_kurt):
        comment = f"Component {i:02d}: Kurtosis z = {z:.2f}"
        if i in focal_inds:
            print(comment + " (Flagged)")
            document_comments['Focality'].append(
                {"comment": comment, "flagged": True})
        else:
            print(comment)
            document_comments['Focality'].append(
                {"comment": comment, "flagged": False})

    # Step 6: Manual inspection and override

    inspect = {
        # "PID 3": [1, 2]
    }

    inspect_components = inspect.get(ppid, list())

    if inspect_components:
        ica.plot_properties(raw, picks=inspect_components)

#     manual_keep = {
#       "PID 3": {1, 2}
#     }

    # Can be filled with manual overrides per ppid
    manual_keep = {}
    manual_exclude = {
        "PID 3": [16, 26],
        "PID 4": [17],
        "PID 5": [9, 15],
        "PID 6": [9],
        "PID 7": [10, 12, 17, 19, 21, 22],
        "PID 8": [2, 3, 10, 12, 15, 25],
        "PID 10": [2, 28],  
        "PID 11": [3, 14],  
        "PID 13": [12, 14],
        "PID 14": [6, 9, 10, 11, 24],
        "PID 15": [6, 7],  
        "PID 16": [15, 12, 4, 22, 26],
        "PID 17": [9, 20],
        "PID 18": [6, 11, 24],
        "PID 19": [5, 6],
        "PID 20": [17, 25],  
        "PID 22": [7, 19],  
        "PID 23": [8, 24],
        "PID 24": [6, 18],
        "PID 25": [3, 9, 10, 13, 15],
        "PID 26": [9, 19],  
        "PID 27": [7, 20, 28],
        "PID 29": [3],
        "PID 30": [11, 19, 22],
        "PID 32": [18, 27, 28],
        "PID 33": [5, 10, 20],
        "PID 35": [3, 11, 16, 17, 22],
        "PID 36": [12, 15, 23],
        "PID 38": [26],
        "PID 41": [6, 7, 10, 20, 22, 27],
        "PID 42": [9, 12, 13, 25, 28],
        "PID 43": [16, 17, 21, 24],
        "PID 45": [14, 18, 21],
        "PID 46": [6, 25],
        "PID 49": [12, 14, 15, 25],
        "PID 50": [7, 8, 12, 18],
        "PID 52": [16, 22],
        "PID 55": [5, 23, 27],
        "PID 58": [7]

    }

    to_keep = manual_keep.get(ppid, set())  # Get components to keep (if any)
    manual_exclude_set = set(manual_exclude.get(ppid, []))
    all_artifacts = (artifact_inds | eog_inds | auto_inds |
                     focal_inds | manual_exclude_set) - to_keep  # Final exclusion list
    ica.exclude = list(all_artifacts)

    # Combine all flagged indices for full documentation
    all_flagged_inds = list(artifact_inds | eog_inds |
                            auto_inds | focal_inds | manual_exclude_set)

    print(f"\n Final components to exclude: {ica.exclude}")

    # Show the ICA component topomap overview interactively
    print("Displaying ICA component overview")
    ica.plot_components(show=True)

    if ica.exclude:
        # Plot properties of excluded components
        ica.plot_properties(raw, picks=ica.exclude)

    # Step 7: Apply ICA and visualize/save before/after comparisons
    raw_before = raw.copy()
    ica.apply(raw)  # Apply ICA exclusion
    raw_after = raw

    # raw plot before ICA
    # print("BEFORE ICA")
    # raw_before.plot(n_channels=30, duration=10, title="Raw Data - Before ICA")

    # raw plot after ICA
    # print("AFTER ICA")
    # raw_after.plot(n_channels=30, duration=10, title="Raw Data - After ICA")

    output_dir = "./ica_comparison_plots"
    os.makedirs(output_dir, exist_ok=True)
    base = str(ppid).replace(" ", "_")

    fig_before = raw_before.plot_psd(fmax=50, show=False)  # PSD before ICA
    fig_before.savefig(os.path.join(
        output_dir, f"{base}_before_psd.png"), dpi=150)
    plt.close(fig_before)

    fig_after = raw_after.plot_psd(fmax=50, show=False)  # PSD after ICA
    fig_after.savefig(os.path.join(
        output_dir, f"{base}_after_psd.png"), dpi=150)
    plt.close(fig_after)

    print(f" ICA complete. PSD plots saved to {output_dir}")

# =============================================================================
#     Document creation and save ica cleaned raw file
# =============================================================================

    # Save a document outlining confidence readings of the raw file
    create_report_document(labels, ica, raw, ppid, all_flagged_inds,
                           document_comments)

    # Save directiory -> ./ICA_cleaned_raws/PID 3
    save_dir = os.path.join("ICA_cleaned_raws", ppid)
    # Make the directory if it does not exist.
    os.makedirs(save_dir, exist_ok=True)

    # Save path -> ./ICA_cleaned_raws/PID 3/PID 3_ica_cleaned_raw.fif
    save_path = os.path.join(save_dir, f"{ppid}_ica_cleaned_raw.fif")
    # Save the raw
    raw.save(save_path, overwrite=True)

    print(f"ICA-cleaned raw saved to: {save_path}")

    # Apply the annotations back to the raw file
    raw.set_annotations(annotations)

# ==================================================
# STEP 8: Save number of retained components to CSV
# ==================================================

    # Total number of components from ICA
    total_components = ica.n_components_

    # Number of excluded components (identified as artifacts)
    excluded_components = len(ica.exclude)

    # Number of retained (non-artifactual) components
    retained_components = total_components - excluded_components

    # Output directory and CSV path
    output_dir = "ica_component_summary"
    os.makedirs(output_dir, exist_ok=True)
    summary_path = os.path.join(output_dir, "retained_ics.csv")

    # Load existing CSV if it exists
    if os.path.exists(summary_path):
        df = pd.read_csv(summary_path)
        # Remove existing entry for this participant if it exists
        df = df[df["participant_id"] != ppid]
    else:
        # Create new DataFrame if file doesn't exist
        df = pd.DataFrame(columns=[
                          "participant_id", "total_components", "excluded_components", "retained_components"])

    # Create new row for the current participant
    new_row = pd.DataFrame([{
        "participant_id": ppid,
        "total_components": total_components,
        "excluded_components": excluded_components,
        "retained_components": retained_components
    }])

    # Append new data and save
    df = pd.concat([df, new_row], ignore_index=True)
    df.to_csv(summary_path, index=False)

    print(
        f"Saved retained IC count for {ppid} to {summary_path} (no duplicates)")

    return raw
